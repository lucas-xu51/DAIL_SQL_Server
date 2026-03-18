from docx import Document
import json

# ================================
# 固定路径（你提供的路径）
# ================================
gold_path = r"C:\Users\grizz\OneDrive\Desktop\COSC448\ideas\model\DAIL-SQL\dataset\process\SPIDER-TEST_SQL_3-SHOT_EUCDISQUESTIONMASK_QA-EXAMPLE_CTX-200_ANS-4096\gold.txt"
pred_path = r"C:\Users\grizz\OneDrive\Desktop\COSC448\ideas\model\DAIL-SQL\dataset\process\SPIDER-TEST_SQL_3-SHOT_EUCDISQUESTIONMASK_QA-EXAMPLE_CTX-200_ANS-4096\RESULTS_MODEL-gpt-4.txt"
dev_json_path = r"C:\Users\grizz\OneDrive\Desktop\COSC448\ideas\model\DAIL-SQL\dataset\spider\dev.json"

# 输出 docx
output_docx = r"C:\Users\grizz\OneDrive\Desktop\COSC448\comparison_output.docx"


# ================================
# 读取 gold SQL
# ================================
def read_gold(path):
    gold_list = []
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                sql = line.split("\t")[0]  # 只有 SQL
                gold_list.append(sql)
    return gold_list


# ================================
# 读取 predict SQL
# ================================
def read_predict(path):
    pred_list = []
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                pred_list.append(line)
    return pred_list


# ================================
# 读取 NL 问题
# ================================
def read_dev_json(path):
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    nl_list = [item["question"] for item in data]
    return nl_list


# ================================
# 主程序：写入 Word
# ================================
gold_list = read_gold(gold_path)
pred_list = read_predict(pred_path)
nl_list = read_dev_json(dev_json_path)

doc = Document()

n = min(len(nl_list), len(gold_list), len(pred_list))

for i in range(n):
    doc.add_heading(f"Example {i+1}", level=2)

    # NL question
    doc.add_paragraph("NL question:")
    doc.add_paragraph(nl_list[i])
    doc.add_paragraph("")

    # Gold
    doc.add_paragraph("Gold SQL:")
    doc.add_paragraph(gold_list[i])
    doc.add_paragraph("")

    # Predict
    doc.add_paragraph("Predict SQL:")
    doc.add_paragraph(pred_list[i])
    doc.add_paragraph("")

    doc.add_paragraph("--------------------------------------------------")

doc.save(output_docx)

print("完成！输出文件：", output_docx)
